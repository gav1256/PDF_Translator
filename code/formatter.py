"""
formatter.py — Markdown + Word (.docx) output.

Three Word layouts:
  * save_word(text, path, output_lang)       — monolingual (target only)
  * save_word(text, path, "bilingual", bilingual=True)
                                             — source + target stacked, direction
                                               chosen per paragraph by script
  * save_word_columns(text, path, src, tgt)  — source + target in a 2-column table
                                               (target | source), copied in spirit
                                               from the RebbeTorahTranslate renderer

All layouts share a בס"ד top-right header and centered page-number footer, and
parse Markdown markers: headings (#, ##, ###), **bold**, *italic*, lists, and
'## Page N' page breaks.
"""

import re
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HEBREW_RE = re.compile(r'[֐-׿]')


@dataclass
class FormatOptions:
    """
    Presentation knobs for the Word renderers.  The defaults reproduce the
    original hard-coded look, so save_word(...) with no options is unchanged.
    """
    body_pt: float = 11.0          # body text size; headings scale off this
    hebrew_font: str = "David"
    latin_font: str = "Calibri"
    line_spacing: float = 1.0      # 1.0 / 1.15 / 1.5 / 2.0
    margin_in: float = 0.7         # all four page margins
    show_bsd: bool = True          # בס"ד header
    bsd_side: str = "right"        # 'right' or 'left'
    show_page_numbers: bool = True # centered PAGE field in the footer
    justify: bool = False          # justify body paragraphs
    swap_columns: bool = False     # columns layout: source | target
    doc_title: str = ""            # written to the .docx Title property

    def font_for(self, is_hebrew: bool) -> str:
        return self.hebrew_font if is_hebrew else self.latin_font

    def heading_pt(self, level: int) -> float:
        """h1/h2/h3 sit 7/5/3 pt above the body (18/16/14 at an 11 pt body)."""
        return max(self.body_pt + 9 - 2 * level, self.body_pt + 1)


DEFAULT_OPTIONS = FormatOptions()


def _apply_para_format(paragraph, opts, justify_ok=False):
    """Line spacing for every paragraph; justification for body text only."""
    if opts.line_spacing and opts.line_spacing != 1.0:
        paragraph.paragraph_format.line_spacing = opts.line_spacing
    if justify_ok and opts.justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def save_markdown(text, output_path):
    """Save raw text to a Markdown file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)
    return output_path


def _is_hebrew_text(s: str) -> bool:
    """True if the string is majority-Hebrew (by letter count)."""
    heb = len(_HEBREW_RE.findall(s))
    latin = len(re.findall(r'[A-Za-z]', s))
    return heb > latin


# In the OOXML schema w:bidi must precede w:jc inside w:pPr.  Appending it
# instead produced out-of-order markup that Word "repaired" by dropping the
# alignment, so the direction is inserted at its proper position.
_AFTER_BIDI = ('w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
               'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap',
               'w:jc', 'w:textDirection', 'w:textAlignment',
               'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
               'w:rPr', 'w:sectPr', 'w:pPrChange')


def _set_bidi(paragraph, rtl: bool):
    """
    Pin a paragraph's direction.  Setting it explicitly to 0 matters as much as
    setting it to 1: a header inside an RTL section inherits that direction
    otherwise, and Word then reads w:jc 'right' as 'end' — the LEFT margin.
    """
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:bidi')):
        pPr.remove(existing)
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1' if rtl else '0')
    pPr.insert_element_before(bidi, *_AFTER_BIDI)


def _set_rtl(paragraph):
    """Set RTL paragraph direction."""
    _set_bidi(paragraph, True)


def _parse_line_type(line):
    """Returns (line_type, level, clean_text)."""
    # Page markers first — otherwise '## Page N' renders as a heading.
    # 'עמוד' is the Hebrew page marker emitted when the target language is Hebrew.
    if re.match(r'^#{1,6}\s+(?:Page|עמוד)\s+\d+\s*$', line):
        return ('page_break', 0, '')
    m = re.match(r'^(#{1,3})\s+(.*)', line)
    if m:
        return ('heading', len(m.group(1)), m.group(2))
    m = re.match(r'^(\s*)-\s+(.*)', line)
    if m:
        return ('bullet', len(m.group(1)) // 2, m.group(2))
    m = re.match(r'^(\s*)\d+[.\)]\s+(.*)', line)
    if m:
        return ('numbered', len(m.group(1)) // 2, m.group(2))
    return ('paragraph', 0, line)


def _is_whole_bold(s: str) -> bool:
    """True if the string is a single **…** span covering the whole thing."""
    s = s.strip()
    return bool(re.fullmatch(r'\*\*[^*]+\*\*', s))


def _add_formatted_runs(paragraph, text, is_hebrew, base_font, base_size,
                        force_bold=False):
    """
    Render markdown into Word runs:
      **bold**, *italic*, ***bold+italic***, and { … } small-print citations
      (rendered at ~0.72x size; italic in the non-Hebrew column for emphasis).
    force_bold bolds every non-citation run (used to mirror whole-paragraph bold
    across the bilingual columns).
    """
    cite_size = Pt(round(base_size.pt * 0.72, 1))

    def _style(run, size, rtl=is_hebrew):
        run.font.name = base_font
        run.font.size = size
        if rtl:
            run.font.rtl = True
            rPr = run._r.get_or_add_rPr()
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), str(int(size.pt * 2)))

    # Split on { citations }, ***/**bold**, *italic* — keep the delimiters.
    parts = re.split(r'(\{[^{}]*\}|\*{3}[^*]+\*{3}|\*{2}[^*]+\*{2}|\*[^*]+\*)', text)
    emitted = False
    for part in parts:
        if not part:
            continue
        is_cite = is_bold = is_italic = False
        if part.startswith('{') and part.endswith('}'):
            is_cite = True
            part = part[1:-1].replace('**', '').strip()
        elif part.startswith('***') and part.endswith('***'):
            is_bold = is_italic = True; part = part[3:-3]
        elif part.startswith('**') and part.endswith('**'):
            is_bold = True; part = part[2:-2]
        elif part.startswith('*') and part.endswith('*'):
            is_italic = True; part = part[1:-1]
        if not part:
            continue
        run = paragraph.add_run(part)
        if (is_bold or force_bold) and not is_cite:
            run.bold = True
        # citations italicised only in the LTR (non-Hebrew) column
        if is_italic or (is_cite and not is_hebrew):
            run.italic = True
        _style(run, cite_size if is_cite else base_size)
        emitted = True

    if not emitted and text:
        _style(paragraph.add_run(text), base_size)


def _setup_document(section_rtl: bool = False, opts=DEFAULT_OPTIONS):
    """Create a Document with margins, בס"ד header, and page-number footer."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(opts.margin_in)
    section.bottom_margin = Inches(opts.margin_in)
    section.left_margin = Inches(opts.margin_in)
    section.right_margin = Inches(opts.margin_in)

    if opts.doc_title:
        doc.core_properties.title = opts.doc_title

    if section_rtl:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

    # Header: בס"ד.  Forced LTR — Word reads w:jc left/right as start/end in an
    # RTL paragraph, so a "right"-aligned run lands against the LEFT margin.
    # The section itself may be RTL (Hebrew targets), hence bidi=0 explicitly
    # rather than merely leaving it unset.  The run's own rtl flag still shapes
    # the Hebrew correctly.
    if opts.show_bsd:
        header_para = section.header.paragraphs[0]
        header_para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if opts.bsd_side == "left"
                                 else WD_ALIGN_PARAGRAPH.RIGHT)
        _set_bidi(header_para, False)
        run = header_para.add_run('בס"ד')
        run.font.name = opts.hebrew_font
        run.font.size = Pt(10)
        run.font.rtl = True

    # Footer: centered PAGE field
    if opts.show_page_numbers:
        footer_para = section.footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = " PAGE "
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        r = footer_para.add_run()
        r._r.append(fld_begin); r._r.append(instr); r._r.append(fld_end)

    return doc, section


def _render_paragraph(doc, line, is_hebrew, base_font, opts=DEFAULT_OPTIONS):
    """Render one markdown line into the document body (shared by mono/stacked)."""
    line_type, level, clean = _parse_line_type(line)

    if line_type == 'page_break':
        if any(p.text.strip() for p in doc.paragraphs):
            doc.add_page_break()
        return

    is_body = False
    if line_type == 'heading':
        p = doc.add_paragraph(style=f'Heading {level}')
        _add_formatted_runs(p, clean, is_hebrew, base_font, Pt(opts.heading_pt(level)))
    elif line_type in ('bullet', 'numbered'):
        p = doc.add_paragraph(style='List Bullet' if line_type == 'bullet' else 'List Number')
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * level)
        _add_formatted_runs(p, clean, is_hebrew, base_font, Pt(opts.body_pt))
    else:
        p = doc.add_paragraph()
        _add_formatted_runs(p, clean, is_hebrew, base_font, Pt(opts.body_pt))
        is_body = True

    if is_hebrew:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl(p)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _apply_para_format(p, opts, justify_ok=is_body)


def save_word(text, output_path, output_lang, bilingual=False, opts=DEFAULT_OPTIONS):
    """
    Monolingual (bilingual=False): one flowing target document, RTL if the
    target language is Hebrew.

    Stacked bilingual (bilingual=True): the interleaved source/target markdown
    rendered paragraph-after-paragraph, with each paragraph's direction chosen
    from its own script (Hebrew paragraphs RTL, others LTR).
    """
    doc_rtl = (not bilingual) and "hebrew" in output_lang.lower()
    doc, _ = _setup_document(section_rtl=doc_rtl, opts=opts)

    target_hebrew = "hebrew" in output_lang.lower()
    for line in text.split('\n'):
        if not line.strip():
            continue
        heb = _is_hebrew_text(line) if bilingual else target_hebrew
        _render_paragraph(doc, line, heb, opts.font_for(heb), opts)

    doc.save(output_path)
    return output_path


def _blocks(text: str) -> list[str]:
    """Split into blocks on blank lines; collapse each block to one line."""
    out = []
    for block in re.split(r'\n\s*\n', text):
        joined = " ".join(l.strip() for l in block.splitlines() if l.strip())
        if joined:
            out.append(joined)
    return out


def _cell_paragraph(cell, first: bool):
    return cell.paragraphs[0] if first else cell.add_paragraph()


def _pair_runs(content_blocks, source_hebrew, target_hebrew):
    """
    Pair a flat list of content blocks into (source, target) tuples.

    When one side is Hebrew, classify each block by script and pair maximal
    same-script runs 1:1 by index — this survives the LLM grouping all-source
    then all-target, or dropping a stray line, without cascade-misaligning the
    rest (mirrors RebbeTorahTranslate's _interleave_grouped).  For pairs where
    neither side is Hebrew, fall back to strict source→target alternation.
    """
    if not (source_hebrew or target_hebrew):
        pairs = []
        for i in range(0, len(content_blocks), 2):
            src = content_blocks[i]
            tgt = content_blocks[i + 1] if i + 1 < len(content_blocks) else ""
            pairs.append((src, tgt))
        return pairs

    def is_source(block):
        return _is_hebrew_text(block) == source_hebrew

    pairs = []
    src_run, tgt_run = [], []

    def flush():
        for i in range(max(len(src_run), len(tgt_run))):
            pairs.append((src_run[i] if i < len(src_run) else "",
                          tgt_run[i] if i < len(tgt_run) else ""))
        src_run.clear()
        tgt_run.clear()

    for block in content_blocks:
        if is_source(block):
            # a new source run begins once we've started seeing targets
            if tgt_run:
                flush()
            src_run.append(block)
        else:
            tgt_run.append(block)
    flush()
    return pairs


def save_word_columns(text, output_path, source_lang, target_lang, opts=DEFAULT_OPTIONS):
    """
    Two-column table: [target | source] (or [source | target] when
    opts.swap_columns).  Headings pair across the columns exactly like body
    text — bold, centred, at heading size — so a title reads across the page
    beside its translation.  '## Page N' markers start a new page rather than
    printing the marker.  Pairing is done by _pair_runs (script-aware).
    """
    source_hebrew = "hebrew" in source_lang.lower()
    target_hebrew = "hebrew" in target_lang.lower()
    doc, section = _setup_document(section_rtl=False, opts=opts)

    col_width = Emu(int((section.page_width - section.left_margin
                         - section.right_margin) / 2))
    table = doc.add_table(rows=0, cols=2)   # borderless by default
    table.columns[0].width = col_width
    table.columns[1].width = col_width

    pending_break = [False]   # the next row opens a new page

    def _fill(cell, text_, is_hebrew, force, size, center):
        p = cell.paragraphs[0]
        _add_formatted_runs(p, text_, is_hebrew, opts.font_for(is_hebrew),
                            size, force_bold=force)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if is_hebrew
                           else WD_ALIGN_PARAGRAPH.LEFT)
        if is_hebrew:
            _set_rtl(p)
        if pending_break[0]:
            p.paragraph_format.page_break_before = True
        # headings keep their centring; only body text may be justified
        _apply_para_format(p, opts, justify_ok=not center)

    def add_row(source, target, size, force, center=False):
        row = table.add_row()
        left, right = row.cells[0], row.cells[1]
        if opts.swap_columns:
            _fill(left, source, source_hebrew, force, size, center)
            _fill(right, target, target_hebrew, force, size, center)
        else:
            _fill(left, target, target_hebrew, force, size, center)
            _fill(right, source, source_hebrew, force, size, center)
        pending_break[0] = False

    # Headings and body text are buffered separately, and each buffer is paired
    # on its own, so a title lines up with its translation rather than with the
    # paragraph that follows it.
    content_buf = []
    heading_buf = []

    def flush_content():
        for src, tgt in _pair_runs(content_buf, source_hebrew, target_hebrew):
            # Mirror whole-paragraph bold: if either side is a single **…**
            # span, bold both cells (keeps a fully-bold quote bold even if the
            # LLM dropped the ** on one column).
            force = _is_whole_bold(src) or _is_whole_bold(tgt)
            add_row(src, tgt, Pt(opts.body_pt), force)
        content_buf.clear()

    def flush_headings():
        if not heading_buf:
            return
        levels = {clean: level for level, clean in heading_buf}
        texts = [clean for _level, clean in heading_buf]
        for src, tgt in _pair_runs(texts, source_hebrew, target_hebrew):
            level = levels.get(src) or levels.get(tgt) or heading_buf[0][0]
            add_row(src, tgt, Pt(opts.heading_pt(level)), True, center=True)
        heading_buf.clear()

    for block in _blocks(text):
        line_type, level, clean = _parse_line_type(block)
        if line_type == 'page_break':
            flush_headings()
            flush_content()
            # The source page number is structure, not something the reader
            # wants printed: break the page instead of writing the marker.
            if len(table.rows):
                pending_break[0] = True
        elif line_type == 'heading':
            flush_content()
            heading_buf.append((level, clean))
        else:
            flush_headings()
            content_buf.append(block)
    flush_headings()
    flush_content()

    doc.save(output_path)
    return output_path
